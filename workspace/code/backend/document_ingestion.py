from typing import List, Dict
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

class DocumentIngestion:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def process_document(self, file_path: str) -> List[Document]:
        """Process a document and return chunks"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(file_path)
        elif file_extension in ['.txt', '.md']:
            return self._process_text(file_path)
        elif file_extension == '.docx':
            return self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path: str) -> List[Document]:
        """Process PDF file"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            
            chunks = self.text_splitter.split_text(text)
            return [Document(page_content=chunk, metadata={"source": file_path}) 
                    for chunk in chunks]
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
    
    def _process_text(self, file_path: str) -> List[Document]:
        """Process text or markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        chunks = self.text_splitter.split_text(text)
        return [Document(page_content=chunk, metadata={"source": file_path}) 
                for chunk in chunks]
    
    def _process_docx(self, file_path: str) -> List[Document]:
        """Process Word document"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            chunks = self.text_splitter.split_text(text)
            return [Document(page_content=chunk, metadata={"source": file_path}) 
                    for chunk in chunks]
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
    
    def process_text_direct(self, text: str, source: str = "direct_input") -> List[Document]:
        """Process text directly without file"""
        chunks = self.text_splitter.split_text(text)
        return [Document(page_content=chunk, metadata={"source": source}) 
                for chunk in chunks]
